from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from ..schemas import DocFormat, Segment
from .base import Translator

# Sentinels: keep inline run boundaries intact across translation. The LLM is
# instructed to preserve ⟦n⟧...⟦/n⟧ markers; when a paragraph has multiple runs
# with different formatting, we join them with these markers. After translation
# we split by markers and re-assign text to the corresponding run.
_SENTINEL_RE = re.compile(r"⟦(\d+)⟧(.*?)⟦/\1⟧", re.DOTALL)


def _paragraph_text_with_sentinels(paragraph) -> tuple[str, list]:
    """Return (joined_text_with_sentinels, list_of_runs)."""
    runs = list(paragraph.runs)
    if len(runs) <= 1:
        return (runs[0].text if runs else ""), runs
    parts: list[str] = []
    for i, r in enumerate(runs, start=1):
        parts.append(f"⟦{i}⟧{r.text}⟦/{i}⟧")
    return "".join(parts), runs


def _split_translated_into_runs(translated: str, n_runs: int) -> list[str]:
    """Given translated text (ideally with sentinels) and expected run count,
    return a list of per-run strings. Falls back gracefully on mismatch."""
    # Pull matches in any order.
    matches = _SENTINEL_RE.findall(translated)
    if not matches:
        # No sentinels survived — put everything in the first run.
        return [translated] + [""] * (n_runs - 1)
    # Map: run_index -> text (1-based from prompt)
    by_idx: dict[int, str] = {}
    for num, text in matches:
        try:
            idx = int(num)
            by_idx[idx] = text
        except ValueError:
            continue
    out: list[str] = []
    for i in range(1, n_runs + 1):
        out.append(by_idx.get(i, ""))
    # If the model merged all into run 1, distribute leftover to empty slots
    # would lose info; better to keep as-is (some runs empty).
    return out


def _iter_paragraphs(doc) -> Iterable:
    """Yield all paragraphs in body + tables + headers + footers."""
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_cell_paragraphs(cell)
    # Header/footer paragraphs in each section
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_cell_paragraphs(cell)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell) -> Iterable:
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                yield from _iter_cell_paragraphs(c)


class DocxTranslator(Translator):
    fmt = DocFormat.DOCX

    def extract(self, src_path: Path) -> list[Segment]:
        from docx import Document

        doc = Document(str(src_path))
        segments: list[Segment] = []
        idx = 0
        for p in _iter_paragraphs(doc):
            joined, runs = _paragraph_text_with_sentinels(p)
            if not joined.strip():
                continue
            segments.append(
                Segment(
                    id=f"p{idx}",
                    text=joined,
                    meta={"n_runs": len(runs), "para_idx": idx},
                )
            )
            idx += 1
        return segments

    def reinsert(self, src_path: Path, segments: list[Segment], out_path: Path) -> Path:
        from docx import Document

        doc = Document(str(src_path))
        translations: dict[int, str] = {}
        for seg in segments:
            translations[seg.meta["para_idx"]] = seg.translated or seg.text

        idx = 0
        for p in _iter_paragraphs(doc):
            joined, runs = _paragraph_text_with_sentinels(p)
            if not joined.strip():
                continue
            tr = translations.get(idx)
            idx += 1
            if tr is None:
                continue
            n = len(runs)
            if n == 0:
                continue
            if n == 1:
                # Strip any sentinels if they slipped in.
                text = _SENTINEL_RE.sub(lambda m: m.group(2), tr)
                runs[0].text = text
            else:
                parts = _split_translated_into_runs(tr, n)
                # If all parts are empty except first (fallback), dump whole text into first run.
                non_empty = sum(1 for p_ in parts if p_)
                if non_empty <= 1:
                    whole = _SENTINEL_RE.sub(lambda m: m.group(2), tr)
                    runs[0].text = whole
                    for r in runs[1:]:
                        r.text = ""
                else:
                    for r, part in zip(runs, parts, strict=True):
                        r.text = part

        doc.save(str(out_path))
        return out_path
